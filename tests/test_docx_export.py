# Mr. Rao -- Copyright (c) 2026 Antonio Andrea Rao.
# SPDX-License-Identifier: AGPL-3.0-or-later
# Software libero: puoi ridistribuirlo e/o modificarlo secondo i termini della
# GNU Affero General Public License pubblicata dalla Free Software Foundation,
# versione 3 o (a tua scelta) successiva. Vedi LICENSE nella radice del repository.
"""L'esportazione in .docx: cosa esce, e soprattutto cosa non ci deve essere.

La prova che conta non e' «il file si apre». E' che il documento consegnato
**non contenga i dati personali**: se un giorno l'esportazione partisse dal
testo originale invece che da quello redatto, il file si aprirebbe lo stesso
e nessuno se ne accorgerebbe.
"""
from __future__ import annotations

import inspect
import io

import pytest

pytest.importorskip("docx", reason="python-docx non installato")

from mr_rao.docx_export import markdown_to_docx  # noqa: E402


def _leggi(dati: bytes):
    from docx import Document

    return Document(io.BytesIO(dati))


def test_il_dato_personale_non_e_nel_documento():
    """Il .docx contiene i segnaposto, non i dati.

    Il controllo guarda **tutto** il testo del file, tabelle comprese: un
    dato lasciato in una cella non e' meno leggibile di uno lasciato in un
    paragrafo.
    """
    md = (
        "Gentile Dott. {{NAME}},\n\nscrivo da {{EMAIL}}, tel. {{PHONE}}.\n\n"
        "| Voce | Valore |\n|---|---|\n| IBAN | {{IBAN}} |\n"
    )
    doc = _leggi(markdown_to_docx(md))

    tutto = "\n".join(p.text for p in doc.paragraphs)
    for tabella in doc.tables:
        for riga in tabella.rows:
            tutto += "\n" + "\n".join(c.text for c in riga.cells)

    for segnaposto in ("{{NAME}}", "{{EMAIL}}", "{{PHONE}}", "{{IBAN}}"):
        assert segnaposto in tutto, f"manca {segnaposto}"
    # e nessuna traccia di dati veri: qui non ce n'erano, ed e' il punto
    assert "@" not in tutto.replace("{{EMAIL}}", "")


def test_il_frontmatter_non_finisce_nel_documento():
    """Le righe YAML servono a chi automatizza, non a chi stampa."""
    md = '---\ngenerator: "Mr. Rao"\nsource: "contratto.pdf"\n---\n\nTesto vero.\n'
    doc = _leggi(markdown_to_docx(md))
    testo = "\n".join(p.text for p in doc.paragraphs)
    assert "Testo vero." in testo
    assert "generator" not in testo and "contratto.pdf" not in testo


def test_le_strutture_diventano_strutture():
    """Intestazioni, elenchi e tabelle non devono appiattirsi in paragrafi."""
    md = (
        "# Titolo\n\n## Sottotitolo\n\n- primo\n- secondo\n\n"
        "1. uno\n2. due\n\n| A | B |\n|---|---|\n| 1 | 2 |\n\n> nota\n"
    )
    doc = _leggi(markdown_to_docx(md))
    stili = [p.style.name for p in doc.paragraphs if p.text.strip()]

    assert "Heading 1" in stili and "Heading 2" in stili
    assert stili.count("List Bullet") == 2
    assert stili.count("List Number") == 2
    assert any("Quote" in s for s in stili)
    assert len(doc.tables) == 1 and doc.tables[0].cell(1, 0).text == "1"


def test_gli_asterischi_del_codice_non_diventano_corsivo():
    """Dentro un blocco di codice il Markdown non si interpreta.

    Senza questo, `codice * con * asterischi` perdeva gli asterischi e usciva
    in corsivo -- cioe' il documento diceva una cosa diversa dall'originale.
    """
    md = "```\nvalore = a * b * c\n```\n"
    doc = _leggi(markdown_to_docx(md))
    testo = "\n".join(p.text for p in doc.paragraphs)
    assert "a * b * c" in testo


def test_il_grassetto_resta_grassetto():
    md = "Questo e' **importante** e questo e' *sottolineato*.\n"
    doc = _leggi(markdown_to_docx(md))
    p = next(p for p in doc.paragraphs if "importante" in p.text)
    assert any(r.bold and r.text == "importante" for r in p.runs)
    assert any(r.italic and r.text == "sottolineato" for r in p.runs)


def test_documento_vuoto_non_esplode():
    """Un Markdown vuoto deve dare un .docx vuoto, non un'eccezione."""
    dati = markdown_to_docx("")
    assert dati and _leggi(dati) is not None


def test_endpoint_rifiuta_il_vuoto(client):
    """Chiedere l'esportazione di niente e' un errore del client, non del server."""
    r = client.post("/api/export/docx", json={"markdown": "   "})
    assert r.status_code == 400


def test_endpoint_restituisce_un_docx(client):
    """Il percorso completo: richiesta, tipo MIME, nome del file."""
    r = client.post(
        "/api/export/docx",
        json={"markdown": "# Titolo\n\nGentile {{NAME}}.\n", "filename": "contratto.pdf"},
    )
    assert r.status_code == 200
    assert "wordprocessingml" in r.headers["Content-Type"]
    assert "contratto.docx" in r.headers.get("Content-Disposition", "")
    # e' davvero uno zip Office, non una pagina d'errore con l'aria giusta
    assert r.data[:2] == b"PK"


def test_il_nome_del_file_non_puo_uscire_dalla_cartella(client):
    """Il nome arriva dal client e finisce in un'intestazione HTTP.

    Non deve poter contenere percorsi ne' ritorni a capo: il primo servirebbe
    a scrivere altrove, il secondo a iniettare altre intestazioni.
    """
    r = client.post(
        "/api/export/docx",
        json={"markdown": "testo", "filename": "../../etc/passwd"},
    )
    assert r.status_code == 200
    disposizione = r.headers.get("Content-Disposition", "")
    assert ".." not in disposizione and "/" not in disposizione.split("filename=")[-1]


def test_le_espressioni_di_riga_ricevono_davvero_una_riga_alla_volta():
    r"""La premessa su cui poggia l'uso di `[ \t]` al posto di `\s`.

    Le espressioni di blocco non attraversano piu' il ritorno a capo. È
    corretto perché il testo viene diviso per righe prima — ma se un domani
    qualcuno passasse un blocco intero a quelle espressioni, smetterebbero di
    riconoscerlo **in silenzio**. Questo test tiene ferma la premessa invece
    di fidarsi che resti vera.
    """
    from mr_rao import docx_export as dx

    sorgente = inspect.getsource(dx.markdown_to_docx)
    assert r'split("\n")' in sorgente, "il testo non viene più diviso per righe"

    # E la conseguenza, dichiarata: con un a capo dentro, non è un titolo.
    assert dx._RE_INTESTAZIONE.match("# titolo")
    assert not dx._RE_INTESTAZIONE.match("#\nnon un titolo")
